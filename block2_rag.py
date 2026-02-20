"""Блок 2: RAG - поиск по базе знаний.
Поддержка гибридного поиска: векторная близость + совпадение ключевых слов (точные термины из запроса)."""
import os
import re
from typing import List, Dict, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
)
from langchain.schema import Document
import config

# Инициализация эмбеддингов
embeddings = HuggingFaceEmbeddings(
    model_name=config.EMBEDDING_MODEL,
    model_kwargs={'device': 'cpu'}
)

# Сплиттер: сначала по абзацам/предложениям, потом по словам, чтобы не резать термины
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=config.CHUNK_SIZE,
    chunk_overlap=config.CHUNK_OVERLAP,
    length_function=len,
    separators=["\n\n", "\n", ". ", ", ", " ", ""],
)

vector_store = None


def _normalize_text_for_indexing(text: str) -> str:
    """Нормализация текста перед разбиением: убираем лишние пробелы/переносы, чтобы не портить чанки."""
    if not text or not text.strip():
        return text
    # Убираем нулевые байты и лишние пробелы/переносы
    t = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    t = re.sub(r"\n{3,}", "\n\n", t)
    t = re.sub(r"[ \t]+", " ", t)
    return t.strip()


def load_knowledge_base():
    """Загружает материалы курса в векторную базу"""
    global vector_store
    
    if not os.path.exists(config.KNOWLEDGE_BASE_PATH):
        os.makedirs(config.KNOWLEDGE_BASE_PATH)
        print(f"Создана папка {config.KNOWLEDGE_BASE_PATH}. Добавьте туда материалы курса (PDF, TXT, MD, DOCX)")
        return None
    
    documents = []
    base_path = os.path.abspath(config.KNOWLEDGE_BASE_PATH)

    # Обход всех файлов в knowledge_base и во вложенных папках (PDF, TXT, MD, DOCX)
    for root, _dirs, files in os.walk(base_path):
        for filename in files:
            filepath = os.path.join(root, filename)
            rel_path = os.path.relpath(filepath, base_path)
            try:
                if filename.lower().endswith('.pdf'):
                    loader = PyPDFLoader(filepath)
                    docs = loader.load()
                    for d in docs:
                        if d.page_content and d.page_content.strip():
                            d.page_content = _normalize_text_for_indexing(d.page_content)
                            if d.page_content:
                                d.metadata["source"] = rel_path
                                documents.append(d)
                elif filename.lower().endswith('.txt'):
                    loader = TextLoader(filepath, encoding='utf-8')
                    docs = loader.load()
                    for d in docs:
                        if d.page_content and d.page_content.strip():
                            d.page_content = _normalize_text_for_indexing(d.page_content)
                            if d.page_content:
                                d.metadata["source"] = rel_path
                                documents.append(d)
                elif filename.lower().endswith('.md'):
                    with open(filepath, 'r', encoding='utf-8') as f:
                        text = _normalize_text_for_indexing(f.read())
                    if text:
                        documents.append(Document(page_content=text, metadata={"source": rel_path}))
                elif filename.lower().endswith('.docx'):
                    from docx import Document as DocxDocument
                    doc = DocxDocument(filepath)
                    text = _normalize_text_for_indexing('\n'.join([p.text for p in doc.paragraphs if p.text]))
                    if text:
                        documents.append(Document(page_content=text, metadata={"source": rel_path}))
            except Exception as e:
                print(f"Ошибка при загрузке {rel_path}: {e}")
                continue
    
    if not documents:
        print(f"Не найдено документов в {config.KNOWLEDGE_BASE_PATH}")
        return None
    
    # Разбивка на чанки
    chunks = text_splitter.split_documents(documents)
    print(f"Загружено {len(documents)} документов, создано {len(chunks)} чанков")
    
    # Создание векторной базы
    os.makedirs(config.VECTOR_DB_PATH, exist_ok=True)
    vector_store = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=config.VECTOR_DB_PATH
    )
    # В Chroma 0.4.x автоматическое сохранение, persist() больше не нужен
    # vector_store.persist()  # Удалено - вызывает предупреждение
    
    print("Векторная база создана и сохранена")
    return vector_store


def _extract_query_terms(query: str) -> List[str]:
    """Извлекает значимые слова из запроса для поиска точных вхождений (кириллица + латиница + цифры)."""
    # Оставляем буквы (в т.ч. кириллица), цифры, дефис; разбиваем по пробелам и знакам
    raw = re.sub(r"[^\w\s\-]", " ", query, flags=re.UNICODE)
    words = [w.strip().lower() for w in raw.split() if w.strip()]
    # Отбрасываем слишком короткие (союзы, предлоги) и длиной не больше 50
    return [w for w in words if 2 <= len(w) <= 50]


def _keyword_score(chunk_text: str, terms: List[str]) -> int:
    """Считает, сколько терминов запроса встречаются в чанке (без учёта регистра)."""
    if not terms:
        return 0
    lower = chunk_text.lower()
    return sum(1 for t in terms if t in lower)


def search_relevant_chunks(query: str, top_k: int = None) -> List[Dict[str, Any]]:
    """
    Гибридный поиск: семантика (эмбеддинги) + совпадение ключевых слов.
    Так находятся и точные термины из базы (ESG, названия и т.д.), и смыслово близкие фрагменты.
    
    Returns:
        List of dicts with keys: content, score, metadata
    """
    global vector_store

    if vector_store is None:
        if os.path.exists(config.VECTOR_DB_PATH):
            try:
                vector_store = Chroma(
                    persist_directory=config.VECTOR_DB_PATH,
                    embedding_function=embeddings,
                )
            except Exception:
                vector_store = load_knowledge_base()
        else:
            vector_store = load_knowledge_base()

    if vector_store is None:
        return []

    top_k = top_k or config.TOP_K
    n_candidates = getattr(config, "TOP_K_CANDIDATES", 16)

    # 1) Берём больше кандидатов по векторной близости
    results = vector_store.similarity_search_with_score(query, k=min(n_candidates, 50))

    terms = _extract_query_terms(query)
    chunks_with_meta = []

    for doc, score in results:
        content = doc.page_content
        # Chroma: меньше score = ближе (L2). Нормализуем в "похожесть": чем меньше distance, тем лучше
        vector_score = float(score)
        kw = _keyword_score(content, terms)
        chunks_with_meta.append({
            "content": content,
            "score": vector_score,
            "metadata": doc.metadata,
            "keyword_hits": kw,
            "terms": terms,
        })

    # 2) Переранжирование: чанки с большим числом совпадений терминов — выше
    def rank_key(c):
        return (c["keyword_hits"], -c["score"])

    chunks_with_meta.sort(key=rank_key, reverse=True)

    # 3) Возвращаем top_k, убираем служебные поля для совместимости
    out = []
    for c in chunks_with_meta[:top_k]:
        out.append({
            "content": c["content"],
            "score": c["score"],
            "metadata": c["metadata"],
        })
    return out


def get_context_from_chunks(chunks: List[Dict[str, Any]]) -> str:
    """Формирует контекст из чанков для промпта"""
    context_parts = []
    for i, chunk in enumerate(chunks, 1):
        source = chunk.get("metadata", {}).get("source", "неизвестный источник")
        context_parts.append(f"[Фрагмент {i} из {source}]\n{chunk['content']}\n")
    
    return "\n".join(context_parts)

